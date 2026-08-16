import logging
import re
import asyncio
from pathlib import Path
from typing import Callable, Set
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from ctranslate2 import Translator
from transformers import AutoTokenizer

logger = logging.getLogger(__name__)


class TranslationService:
    _instance = None
    _translator = None
    _tokenizer = None

    def __new__(cls):
        if cls._instance is None:
            cls._instance = super().__new__(cls)
        return cls._instance

    def __init__(self):
        self.progress_callbacks = {}
        self.cancelled_requests: Set[str] = set()
        self.main_loop = None
        self.temp_dir = Path("./tmp")
        self.temp_dir.mkdir(parents=True, exist_ok=True)

    def load_model(self, model_path: Path, device: str = "cpu"):
        if self._translator is None:
            logger.info(f"Loading model and tokenizer from {model_path}")
            self._translator = Translator(str(model_path), device=device, compute_type="int8")
            self._tokenizer = AutoTokenizer.from_pretrained(str(model_path), src_lang="eng_Latn")
            logger.info("Model and tokenizer loaded successfully")
        return self._translator

    def cancel_request(self, request_id: str):
        """Отмена операции по request_id"""
        self.cancelled_requests.add(request_id)
        logger.info(f"Cancelled request: {request_id}")

    def is_cancelled(self, request_id: str) -> bool:
        return request_id in self.cancelled_requests

    def set_progress_callback(self, request_id: str, callback: Callable):
        """Установка callback для прогресса"""
        self.progress_callbacks[request_id] = callback
        logger.info(f"Callback set for {request_id}")

    def remove_progress_callback(self, request_id: str):
        """Удаление callback"""
        if request_id in self.progress_callbacks:
            del self.progress_callbacks[request_id]
        if request_id in self.cancelled_requests:
            self.cancelled_requests.remove(request_id)
        logger.info(f"Callback removed for {request_id}")

    async def update_progress(self, request_id: str, progress: int, status: str = "", preview: dict = None):
        """Обновление прогресса"""
        logger.info(f"Progress {request_id}: {progress}% - {status}")

        if self.is_cancelled(request_id):
            raise Exception("Operation cancelled")

        if request_id in self.progress_callbacks:
            callback = self.progress_callbacks[request_id]
            try:
                if asyncio.iscoroutinefunction(callback):
                    await callback(progress, status, preview)
                else:
                    callback(progress, status, preview)
                logger.info(f"Callback executed for {request_id}")
            except Exception as e:
                logger.error(f"Callback error: {e}")
        else:
            logger.warning(f"No callback found for {request_id}")

    def _clean_text(self, text: str) -> str:
        """Умная очистка: убирает мусор и спасает токенизатор от <unk>"""
        if not text:
            return text

        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', text)

        replacements = {
            '¬': '', '\u200b': '', '\u200c': '', '\u200d': '', '\ufeff': ''
        }
        for old, new in replacements.items():
            text = text.replace(old, new)

        text = re.sub(r'(\d)\s*[—–−]\s*(\d)', r'\1-\2', text)
        text = re.sub(r'[ \t]+', ' ', text)

        return text.strip()

    def _post_process_typography(self, text: str) -> str:
        if not text:
            return text
        text = re.sub(r' - ', ' — ', text)
        return text

    def _smart_chunk(self, text: str, max_chars: int = 800) -> list[str]:
        sentences = re.split(r'(?<=[.!?…])\s+', text)
        chunks = []
        current_chunk = ""

        for sentence in sentences:
            if not sentence.strip():
                continue

            if len(sentence) > max_chars:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                    current_chunk = ""
                sub_chunks = re.split(r'(?<=[,;])\s+', sentence)
                temp_chunk = ""
                for sub in sub_chunks:
                    if len(temp_chunk) + len(sub) > max_chars:
                        chunks.append(temp_chunk.strip())
                        temp_chunk = sub + " "
                    else:
                        temp_chunk += sub + " "
                if temp_chunk:
                    current_chunk = temp_chunk
                continue

            if len(current_chunk) + len(sentence) + 1 <= max_chars:
                current_chunk += sentence + " "
            else:
                chunks.append(current_chunk.strip())
                current_chunk = sentence + " "

        if current_chunk:
            chunks.append(current_chunk.strip())

        return chunks

    def _translate_chunk(self, chunk: str, src_lang: str, tgt_lang: str, num_beams: int = 2) -> str:
        if not chunk.strip():
            return ""

        try:
            self._tokenizer.src_lang = src_lang
            input_tokens = self._tokenizer.encode(chunk)
            tokens = self._tokenizer.convert_ids_to_tokens(input_tokens)

            if not tokens:
                return ""

            results = self._translator.translate_batch(
                [tokens],
                target_prefix=[[tgt_lang]],
                max_decoding_length=1024,
                beam_size=num_beams
            )

            if not results or not results[0].hypotheses:
                return chunk

            hypothesis = results[0].hypotheses[0]

            if hypothesis and hypothesis[0] == tgt_lang:
                hypothesis = hypothesis[1:]

            output_ids = self._tokenizer.convert_tokens_to_ids(hypothesis)
            result_text = self._tokenizer.decode(output_ids)

            return result_text.strip()

        except Exception as e:
            logger.error(f"Chunk translation error: {e}")
            return chunk

    def translate_text(self, text: str, src_lang: str = "eng_Latn", tgt_lang: str = "rus_Cyrl", num_beams: int = 2) -> str:
        if not text or not text.strip():
            return text

        if self._translator is None or self._tokenizer is None:
            logger.error("Model is not loaded! Call load_model() first.")
            return text

        lines = text.splitlines()
        translated_lines = []

        for line in lines:
            if not line.strip():
                translated_lines.append("")
                continue

            clean_line = self._clean_text(line)
            chunks = self._smart_chunk(clean_line, max_chars=400)

            translated_chunks = []
            for chunk in chunks:
                translated = self._translate_chunk(chunk, src_lang, tgt_lang, num_beams=num_beams)
                translated_chunks.append(translated)

            translated_lines.append(" ".join(translated_chunks))

        final_text = "\n".join(translated_lines)
        return self._post_process_typography(final_text)

    def _preserve_runs_formatting(self, paragraph: Paragraph, translated_text: str) -> None:
        if not paragraph.runs:
            paragraph.text = translated_text
            return

        if len(paragraph.runs) == 1:
            paragraph.runs[0].text = translated_text
            return

        first_run = paragraph.runs[0]
        original_style = first_run.style
        original_bold = first_run.bold
        original_italic = first_run.italic
        original_underline = first_run.underline
        original_font = first_run.font

        for run in paragraph.runs:
            run.text = ""

        new_run = paragraph.runs[0]
        new_run.text = translated_text

        if original_style:
            new_run.style = original_style
        if original_bold is not None:
            new_run.bold = original_bold
        if original_italic is not None:
            new_run.italic = original_italic
        if original_underline is not None:
            new_run.underline = original_underline
        if original_font:
            new_run.font.name = original_font.name
            new_run.font.size = original_font.size

    def translate_paragraph(self, paragraph: Paragraph, num_beams: int = 2) -> None:
        original_text = paragraph.text
        if not original_text or not original_text.strip():
            return

        translated_text = self.translate_text(original_text, num_beams=num_beams)
        self._preserve_runs_formatting(paragraph, translated_text)

    def translate_table(self, table: Table, num_beams: int = 2) -> None:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    self.translate_paragraph(paragraph, num_beams=num_beams)
                for nested_table in cell.tables:
                    self.translate_table(nested_table, num_beams=num_beams)

    # СИНХРОННЫЙ МЕТОД (без прогресса, для тестов)
    def translate_docx_sync(self, input_path: Path, output_path: Path, num_beams: int = 2) -> Path:
        """Синхронный перевод DOCX без прогресса"""
        if num_beams not in (1, 2, 4):
            num_beams = 2

        doc = Document(input_path)

        for paragraph in doc.paragraphs:
            self.translate_paragraph(paragraph, num_beams=num_beams)

        for table in doc.tables:
            self.translate_table(table, num_beams=num_beams)

        doc.save(output_path)
        logger.info(f"Translated document saved to {output_path} with num_beams={num_beams}")
        return output_path

    # АСИНХРОННЫЙ МЕТОД (с SSE прогрессом)
    async def translate_docx(
        self,
        input_path: Path,
        output_path: Path,
        num_beams: int = 2,
        request_id: str = None
    ) -> Path:
        """Асинхронный перевод с SSE прогрессом и превью"""
        if num_beams not in (1, 2, 4):
            num_beams = 2

        if self.main_loop is None:
            self.main_loop = asyncio.get_running_loop()

        if request_id:
            await self.update_progress(request_id, 5, "Загрузка документа...")

        doc = Document(input_path)
        total_paragraphs = len(doc.paragraphs)

        if request_id:
            await self.update_progress(request_id, 10, f"Начинаем перевод... ({total_paragraphs} абзацев)")

        for idx, paragraph in enumerate(doc.paragraphs):
            if request_id and self.is_cancelled(request_id):
                raise Exception("Operation cancelled")

            original_text = paragraph.text
            if original_text and original_text.strip():
                translated_text = await asyncio.to_thread(
                    self.translate_text,
                    original_text,
                    "eng_Latn",
                    "rus_Cyrl",
                    num_beams
                )
                self._preserve_runs_formatting(paragraph, translated_text)

            progress = 10 + int(((idx + 1) / total_paragraphs) * 80)

            if request_id:
                preview_data = {
                    'paragraph_index': idx,
                    'original': original_text[:300] if original_text else "",
                    'translated': paragraph.text[:300] if paragraph.text else "",
                }
                await self.update_progress(
                    request_id,
                    progress,
                    f"Перевод абзаца {idx + 1}/{total_paragraphs}",
                    preview_data
                )

        total_tables = len(doc.tables)
        if total_tables > 0:
            for table_idx, table in enumerate(doc.tables):
                if request_id and self.is_cancelled(request_id):
                    raise Exception("Operation cancelled")

                await asyncio.to_thread(
                    self.translate_table,
                    table,
                    num_beams
                )

                progress = 90 + int(((table_idx + 1) / total_tables) * 5)
                if request_id:
                    await self.update_progress(
                        request_id,
                        progress,
                        f"Перевод таблицы {table_idx + 1}/{total_tables}",
                        {'table_index': table_idx, 'is_table': True}
                    )

        if request_id:
            await self.update_progress(request_id, 95, "Сохранение документа...")

        await asyncio.to_thread(doc.save, str(output_path))

        if request_id:
            await self.update_progress(request_id, 100, "Готово!")

        logger.info(f"Translated document saved to {output_path} with num_beams={num_beams}")
        return output_path

    def cleanup(self, file_path: Path):
        if file_path and file_path.exists():
            try:
                file_path.unlink()
                logger.info(f"Deleted: {file_path}")
            except Exception as e:
                logger.error(f"Error deleting {file_path}: {e}")

    def unload(self):
        self._translator = None
        self._tokenizer = None
        self.progress_callbacks.clear()
        self.cancelled_requests.clear()
        logger.info("Model unloaded")


translation_service = TranslationService()
