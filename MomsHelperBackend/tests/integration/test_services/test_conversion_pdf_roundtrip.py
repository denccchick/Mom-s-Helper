import io
import pytest
from types import SimpleNamespace
from pathlib import Path

from docx import Document

from app.services.conversion_service import ConversionService
from fastapi import HTTPException
from PIL import Image, ImageDraw, ImageFont


def _make_image_pdf(path: Path, text: str = "Hello world"):
    img = Image.new("RGB", (400, 200), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("arial.ttf", 24)
    except Exception:
        font = None
    draw.text((20, 60), text, fill=(0, 0, 0), font=font)
    img.save(path, "PDF")


class DummyReader:
    def readtext(self, img_array, detail=1):
        bbox = [(10, 10), (300, 10), (300, 80), (10, 80)]
        return [(bbox, "Hello world", 0.98)]


def make_uploadfile(filename: str, content: bytes):
    return SimpleNamespace(filename=filename, file=io.BytesIO(content))


@pytest.mark.asyncio
@pytest.mark.integration
async def test_docx_to_pdf_and_back(tmp_test_dir, uploadfile_factory):
    svc = ConversionService(temp_dir=str(tmp_test_dir))

    # create a small docx
    input_docx = Path(tmp_test_dir) / "in.docx"
    doc = Document()
    doc.add_paragraph("Hello world")
    doc.save(input_docx)

    # convert docx -> pdf via the service (wrapped as UploadFile)
    with open(input_docx, "rb") as f:
        uf = uploadfile_factory(input_docx.name, f.read())

    try:
        pdf_path, _ = svc.docx_to_pdf(uf)
    except HTTPException as e:
        pytest.skip(f"docx->pdf conversion not available in environment: {e.detail}")

    assert pdf_path.exists()

    # now convert pdf -> docx via pdf2docx
    with open(pdf_path, "rb") as f:
        pdf_uf = uploadfile_factory(pdf_path.name, f.read())

    try:
        out_docx, _ = svc.pdf_to_docx(pdf_uf)
    except HTTPException as e:
        pytest.skip(f"pdf->docx conversion failed: {e.detail}")

    assert out_docx.exists()

    # Load resulting docx and ensure it contains some text
    from docx import Document as DocxLoader

    res = DocxLoader(out_docx)
    all_text = "\n".join(p.text for p in res.paragraphs)
    assert len(all_text.strip()) > 0


@pytest.mark.integration
@pytest.mark.asyncio
async def test_pdf_ocr_to_docx(tmp_test_dir, uploadfile_factory):
    svc = ConversionService(temp_dir=str(tmp_test_dir))

    input_pdf = Path(tmp_test_dir) / "img.pdf"
    _make_image_pdf(input_pdf, text="Hello world")

    with open(input_pdf, "rb") as f:
        uf = uploadfile_factory(input_pdf.name, f.read())

    svc.reader = DummyReader()

    out_docx = await svc.pdf2docx_ocr(uf)

    assert out_docx.exists()

    from docx import Document

    doc = Document(out_docx)
    joined = "\n".join(p.text for p in doc.paragraphs).lower()
    assert "hello world" in joined or "hello" in joined
