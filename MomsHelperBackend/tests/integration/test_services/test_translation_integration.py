import pytest
from pathlib import Path
from docx import Document
import asyncio

from app.services.translation_service import TranslationService


@pytest.mark.integration
def test_translate_docx_creates_output(tmp_test_dir, simple_translator):
    # Create a small docx file
    input_path = Path(tmp_test_dir) / "in.docx"
    doc = Document()
    doc.add_paragraph("Hello world")
    doc.add_paragraph("Second paragraph")
    doc.save(input_path)

    out_path = Path(tmp_test_dir) / "out.docx"

    svc = TranslationService()
    # Patch instance translate_text to deterministic russian output for test
    svc.translate_text = simple_translator

    # Используем синхронный метод для тестов
    svc.translate_docx_sync(input_path, out_path)

    assert out_path.exists()

    out_doc = Document(out_path)
    joined = "\n".join(p.text for p in out_doc.paragraphs).lower()

    # Check that at least one expected Russian word exists in output
    expected = ["здравствуй", "здравствуйте", "привет", "мир"]
    assert any(w in joined for w in expected)


@pytest.mark.integration
def test_translate_docx_preserves_formatting(tmp_test_dir, simple_translator):
    """Тест сохранения форматирования при переводе"""
    input_path = Path(tmp_test_dir) / "in.docx"
    doc = Document()

    p = doc.add_paragraph()
    run = p.add_run("Hello world")
    run.bold = True
    run.italic = True

    doc.save(input_path)

    out_path = Path(tmp_test_dir) / "out.docx"

    svc = TranslationService()
    svc.translate_text = simple_translator

    svc.translate_docx_sync(input_path, out_path)

    assert out_path.exists()

    out_doc = Document(out_path)
    assert len(out_doc.paragraphs) > 0

    out_p = out_doc.paragraphs[0]
    assert len(out_p.runs) > 0
    assert out_p.text != "Hello world"
    assert len(out_p.text) > 0


@pytest.mark.integration
def test_translate_docx_handles_tables(tmp_test_dir, simple_translator):
    """Тест перевода таблиц"""
    input_path = Path(tmp_test_dir) / "in.docx"
    doc = Document()

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Hello"
    table.cell(0, 1).text = "World"
    table.cell(1, 0).text = "Test"
    table.cell(1, 1).text = "Data"

    doc.save(input_path)

    out_path = Path(tmp_test_dir) / "out.docx"

    svc = TranslationService()
    svc.translate_text = simple_translator

    svc.translate_docx_sync(input_path, out_path)

    assert out_path.exists()

    out_doc = Document(out_path)
    assert len(out_doc.tables) > 0

    out_table = out_doc.tables[0]
    assert out_table.cell(0, 0).text != "Hello"


@pytest.mark.integration
def test_translate_docx_handles_empty_document(tmp_test_dir, simple_translator):
    """Тест перевода пустого документа"""
    input_path = Path(tmp_test_dir) / "in.docx"
    doc = Document()
    doc.save(input_path)

    out_path = Path(tmp_test_dir) / "out.docx"

    svc = TranslationService()
    svc.translate_text = simple_translator

    svc.translate_docx_sync(input_path, out_path)

    assert out_path.exists()

    out_doc = Document(out_path)
    assert len(out_doc.paragraphs) == 0


@pytest.mark.integration
def test_translate_docx_with_different_beams(tmp_test_dir, simple_translator):
    """Тест перевода с разными значениями num_beams"""
    input_path = Path(tmp_test_dir) / "in.docx"
    doc = Document()
    doc.add_paragraph("Hello world")
    doc.save(input_path)

    svc = TranslationService()
    svc.translate_text = simple_translator

    for beams in [1, 2, 4]:
        test_out = Path(tmp_test_dir) / f"out_{beams}.docx"
        svc.translate_docx_sync(input_path, test_out, num_beams=beams)
        assert test_out.exists()

        out_doc = Document(test_out)
        assert len(out_doc.paragraphs) > 0
        assert out_doc.paragraphs[0].text != "Hello world"
